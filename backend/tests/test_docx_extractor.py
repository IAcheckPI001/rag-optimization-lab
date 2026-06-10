from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile

import pytest

from docx import Document

from app.providers.extraction.docx_extractor import DocxExtractor
from app.providers.extraction.errors import (
    ExtractionNoContentError,
    ExtractionParsingError,
    ExtractionSourceTypeMismatchError,
)
from app.schemas.document import DocumentContentType
from app.schemas.extraction import ExtractionInput, ExtractionResult
from app.schemas.source import SourceType


def docx_bytes(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def extraction_input(content_bytes: bytes, source_type: SourceType = SourceType.docx):
    return ExtractionInput(
        source_id="src_001",
        document_id="doc_001",
        source_type=source_type,
        source_uri="storage://sources/src_001/raw.docx",
        original_filename="policy.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        content_bytes=content_bytes,
    )


def extract(document: Document) -> ExtractionResult:
    return DocxExtractor(extractor_version="test-version").extract(
        extraction_input(docx_bytes(document))
    )


def test_docx_extractor_extracts_single_paragraph() -> None:
    document = Document()
    document.add_paragraph("Leave policy allows 12 paid days.")

    result = extract(document)
    unit = result.units[0]

    assert result.extractor_name == "python-docx"
    assert result.extractor_version == "test-version"
    assert unit.content == "Leave policy allows 12 paid days."
    assert unit.content_type is DocumentContentType.paragraph
    assert unit.unit_index == 0
    assert unit.raw_unit_id == "raw:doc_001:000000"
    assert unit.page_start is None
    assert unit.page_end is None
    assert unit.section is None
    assert unit.heading_path == []
    assert unit.extra_metadata["block_type"] == "paragraph"
    assert unit.extra_metadata["block_index"] == 0
    assert unit.extra_metadata["paragraph_index"] == 0
    assert unit.character_count == len(unit.content)
    assert result.stats.total_units == 1
    assert result.stats.skipped_items == 0
    assert result.stats.extra_metadata["total_body_items"] == 1
    assert result.stats.extra_metadata["paragraph_count"] == 1


def test_docx_extractor_preserves_paragraph_and_table_body_order() -> None:
    document = Document()
    document.add_paragraph("Paragraph A")
    table_a = document.add_table(rows=1, cols=2)
    table_a.cell(0, 0).text = "A1"
    table_a.cell(0, 1).text = "A2"
    document.add_paragraph("Paragraph B")
    table_b = document.add_table(rows=1, cols=2)
    table_b.cell(0, 0).text = "B1"
    table_b.cell(0, 1).text = "B2"

    result = extract(document)

    assert [unit.content for unit in result.units] == [
        "Paragraph A",
        "A1\tA2",
        "Paragraph B",
        "B1\tB2",
    ]
    assert [unit.unit_index for unit in result.units] == [0, 1, 2, 3]
    assert [unit.raw_unit_id for unit in result.units] == [
        "raw:doc_001:000000",
        "raw:doc_001:000001",
        "raw:doc_001:000002",
        "raw:doc_001:000003",
    ]
    assert [unit.extra_metadata["block_index"] for unit in result.units] == [
        0,
        1,
        2,
        3,
    ]
    assert result.stats.extra_metadata["total_body_items"] == 4
    assert result.stats.extra_metadata["paragraph_count"] == 2
    assert result.stats.extra_metadata["table_count"] == 2


def test_docx_extractor_skips_blank_paragraph_without_warning() -> None:
    document = Document()
    document.add_paragraph("First")
    document.add_paragraph("   ")
    document.add_paragraph("Second")

    result = extract(document)

    assert [unit.content for unit in result.units] == ["First", "Second"]
    assert [unit.unit_index for unit in result.units] == [0, 1]
    assert result.warnings == []
    assert result.stats.skipped_items == 1
    assert result.stats.extra_metadata["total_body_items"] == 3
    assert result.stats.extra_metadata["paragraph_count"] == 3


def test_docx_extractor_builds_heading_paths_and_resets_lower_levels() -> None:
    document = Document()
    document.add_paragraph("Human Resources", style="Heading 1")
    document.add_paragraph("Leave Policy", style="Heading 2")
    document.add_paragraph("Annual Leave", style="Heading 3")
    document.add_paragraph("Employees receive annual leave.")
    document.add_paragraph("Salary Policy", style="Heading 2")
    document.add_paragraph("Salary paragraph.")

    result = extract(document)

    assert result.units[0].heading_path == ["Human Resources"]
    assert result.units[0].section == "Human Resources"
    assert result.units[0].extra_metadata["block_type"] == "heading"
    assert result.units[0].extra_metadata["heading_level"] == 1
    assert result.units[1].heading_path == ["Human Resources", "Leave Policy"]
    assert result.units[2].heading_path == [
        "Human Resources",
        "Leave Policy",
        "Annual Leave",
    ]
    assert result.units[3].heading_path == [
        "Human Resources",
        "Leave Policy",
        "Annual Leave",
    ]
    assert result.units[3].section == "Annual Leave"
    assert result.units[4].heading_path == ["Human Resources", "Salary Policy"]
    assert result.units[5].heading_path == ["Human Resources", "Salary Policy"]
    assert result.stats.extra_metadata["heading_count"] == 4


def test_docx_extractor_does_not_update_heading_state_for_blank_heading() -> None:
    document = Document()
    document.add_paragraph("   ", style="Heading 1")
    document.add_paragraph("Paragraph after blank heading.")

    result = extract(document)

    assert len(result.units) == 1
    assert result.units[0].content == "Paragraph after blank heading."
    assert result.units[0].heading_path == []
    assert result.units[0].section is None
    assert result.stats.skipped_items == 1
    assert result.stats.extra_metadata["paragraph_count"] == 2
    assert result.stats.extra_metadata["heading_count"] == 0


def test_docx_extractor_supports_heading_level_gaps_without_synthetic_headings() -> None:
    document = Document()
    document.add_paragraph("Top", style="Heading 1")
    document.add_paragraph("Deep", style="Heading 3")
    document.add_paragraph("Body")

    result = extract(document)

    assert result.units[1].heading_path == ["Top", "Deep"]
    assert result.units[2].heading_path == ["Top", "Deep"]


def test_docx_extractor_table_inherits_heading_path() -> None:
    document = Document()
    document.add_paragraph("Leave Policy", style="Heading 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Leave type"
    table.cell(0, 1).text = "Days"
    table.cell(1, 0).text = "Annual leave"
    table.cell(1, 1).text = "12"

    result = extract(document)
    table_unit = result.units[1]

    assert table_unit.content == "Leave type\tDays\nAnnual leave\t12"
    assert table_unit.content_type is DocumentContentType.table
    assert table_unit.heading_path == ["Leave Policy"]
    assert table_unit.section == "Leave Policy"
    assert table_unit.extra_metadata["block_type"] == "table"
    assert table_unit.extra_metadata["row_count"] == 2
    assert table_unit.extra_metadata["column_count"] == 2
    assert table_unit.extra_metadata["row_column_counts"] == [2, 2]
    assert table_unit.extra_metadata["serialization_format"] == "tsv_escaped_v1"


def test_docx_extractor_escapes_table_cell_text_deterministically() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "A\tB"
    table.cell(0, 1).text = "Line 1\nLine 2"
    table.cell(0, 2).text = r"C:\Docs"

    result = extract(document)

    assert result.units[0].content == r"A\tB	Line 1\nLine 2	C:\\Docs"


def test_docx_extractor_skips_blank_table_based_on_raw_cell_text() -> None:
    document = Document()
    document.add_paragraph("Before")
    document.add_table(rows=2, cols=2)
    document.add_paragraph("After")

    result = extract(document)

    assert [unit.content for unit in result.units] == ["Before", "After"]
    assert result.stats.skipped_items == 1
    assert result.stats.extra_metadata["table_count"] == 1
    assert result.stats.extra_metadata["total_body_items"] == 3


def test_docx_extractor_ignores_structural_section_properties_for_counters() -> None:
    document = Document()
    document.add_paragraph("Only content")

    result = extract(document)

    assert result.units[0].extra_metadata["block_index"] == 0
    assert result.stats.extra_metadata["total_body_items"] == 1
    assert result.stats.skipped_items == 0
    assert result.stats.extra_metadata["unsupported_item_count"] == 0


def test_docx_extractor_is_deterministic_except_extracted_at() -> None:
    document = Document()
    document.add_paragraph("Policy", style="Heading 1")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Type"
    table.cell(0, 1).text = "Days"
    content = docx_bytes(document)

    first = DocxExtractor(extractor_version="test-version").extract(
        extraction_input(content)
    )
    second = DocxExtractor(extractor_version="test-version").extract(
        extraction_input(content)
    )

    comparable_first = [
        (
            unit.content,
            unit.content_type,
            unit.unit_index,
            unit.raw_unit_id,
            unit.page_start,
            unit.page_end,
            unit.section,
            unit.heading_path,
            unit.extra_metadata,
        )
        for unit in first.units
    ]
    comparable_second = [
        (
            unit.content,
            unit.content_type,
            unit.unit_index,
            unit.raw_unit_id,
            unit.page_start,
            unit.page_end,
            unit.section,
            unit.heading_path,
            unit.extra_metadata,
        )
        for unit in second.units
    ]

    assert comparable_first == comparable_second
    assert len({unit.extracted_at for unit in first.units}) == 1


def test_docx_extractor_rejects_source_type_mismatch() -> None:
    document = Document()
    document.add_paragraph("PDF mismatch")

    with pytest.raises(ExtractionSourceTypeMismatchError) as exc_info:
        DocxExtractor().extract(extraction_input(docx_bytes(document), SourceType.pdf))

    assert exc_info.value.error_code == "extraction_source_type_mismatch"
    assert "content_bytes" not in exc_info.value.details


@pytest.mark.parametrize("content_bytes", [b"not a docx", b"PK\x03\x04broken"])
def test_docx_extractor_raises_parsing_error_for_invalid_docx(
    content_bytes: bytes,
) -> None:
    with pytest.raises(ExtractionParsingError) as exc_info:
        DocxExtractor().extract(extraction_input(content_bytes))

    assert exc_info.value.error_code == "extraction_parsing_failed"
    assert "content_bytes" not in exc_info.value.details


def test_docx_extractor_raises_parsing_error_for_corrupted_docx_package() -> None:
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "not xml")

    with pytest.raises(ExtractionParsingError):
        DocxExtractor().extract(extraction_input(stream.getvalue()))


def test_docx_extractor_raises_no_content_for_empty_docx() -> None:
    with pytest.raises(ExtractionNoContentError) as exc_info:
        DocxExtractor().extract(extraction_input(docx_bytes(Document())))

    assert exc_info.value.error_code == "extraction_no_content"
    assert "content_bytes" not in exc_info.value.details


def test_docx_extractor_raises_no_content_for_only_blank_paragraphs() -> None:
    document = Document()
    document.add_paragraph("   ")
    document.add_paragraph("\t")

    with pytest.raises(ExtractionNoContentError):
        DocxExtractor().extract(extraction_input(docx_bytes(document)))
